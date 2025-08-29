"""
pdf_word_to_text.py

Dependencies (install with pip if you don't have them):
    pip install pdfplumber PyPDF2 python-docx
"""

from pathlib import Path
from typing import Optional, Sequence


from pypdf import PdfReader
from fastapi import UploadFile

async def pdf_to_text(file: UploadFile) -> str:
    """
    Extract text from a PDF UploadFile using pypdf.
    """
    text = ""

    # Reset file pointer to beginning (in case it was read elsewhere)
    await file.seek(0)

    reader = PdfReader(file.file)
    for page in reader.pages:
        text += page.extract_text() or ""

    return text


def docx_to_text(docx_path: str | Path,
                 output_path: Optional[str | Path] = None) -> str:
    """
    Extract text from a .docx Word document.

    Args:
        docx_path: path to the .docx file.
        output_path: optional path to write the extracted text (.txt).

    Returns:
        The extracted text as a single string.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f".docx not found: {docx_path}")

    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to read .docx files. Install with `pip install python-docx`.\n"
            f"Original error: {e}"
        )

    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    # Also try to extract text from tables (if present)
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                table_texts.append("\t".join(cells))

    combined_parts = []
    if paragraphs:
        combined_parts.append("\n".join(paragraphs))
    if table_texts:
        combined_parts.append("\n".join(table_texts))

    full_text = "\n\n".join(combined_parts).strip()

    if output_path:
        Path(output_path).write_text(full_text, encoding="utf-8")

    return full_text


# Example usage
if __name__ == "__main__":
    # Example 1: Extract all pages from PDF and print first 500 chars
    pdf_text = pdf_to_text("example.pdf", output_path="example_pdf.txt")
    print("PDF extracted (first 500 chars):")
    print(pdf_text[:500])

    # Example 2: Extract Word docx content
    word_text = docx_to_text("example.docx", output_path="example_docx.txt")
    print("\nDOCX extracted (first 500 chars):")
    print(word_text[:500])
